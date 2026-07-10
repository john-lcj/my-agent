"""docx_writer skill:把 Markdown/文本生成 Word(.docx)。"""
from __future__ import annotations

import os

from core.types import CapabilityResult
from governance.workspace import resolve_path

RISK = "WRITE"

SCHEMA = {
    "type": "object",
    "properties": {
        "path": {"type": "string", "description": "Output .docx path"},
        "markdown": {"type": "string", "description": "Body text, Markdown or plain text"},
        "title": {"type": "string", "description": "Document title; optional"},
    },
    "required": ["path", "markdown"],
}


def _flush_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        pass
    for r in rows:
        cells = table.add_row().cells
        for i in range(ncol):
            cells[i].text = r[i] if i < len(r) else ""


async def run(args: dict, ctx) -> CapabilityResult:
    path, error = resolve_path(str(args.get("path", "")).strip())
    md = str(args.get("markdown", ""))
    if error:
        return CapabilityResult(ok=False, error=error)
    if not path or not md.strip():
        return CapabilityResult(ok=False, error="缺少 path 或 markdown")
    if not path.endswith(".docx"):
        path += ".docx"
    try:
        from docx import Document
    except Exception as e:
        return CapabilityResult(ok=False, error=f"需要 python-docx:{e}")

    doc = Document()
    title = str(args.get("title", "")).strip()
    if title:
        doc.add_heading(title, level=0)

    n_para = n_tbl = 0
    table_buf: list[list[str]] = []
    for raw in md.splitlines():
        line = raw.rstrip()
        is_table_row = line.strip().startswith("|") and line.strip().endswith("|")
        if is_table_row:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if set("".join(cells)) <= set("-: "):  # 分隔行 |---|---|
                continue
            table_buf.append(cells)
            continue
        if table_buf:
            _flush_table(doc, table_buf); n_tbl += 1; table_buf = []
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:], style="List Bullet")
        else:
            doc.add_paragraph(s)
        n_para += 1
    if table_buf:
        _flush_table(doc, table_buf); n_tbl += 1

    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        doc.save(path)
    except Exception as e:
        return CapabilityResult(ok=False, error=f"写入失败:{e}")
    return CapabilityResult(ok=True, output=f"已生成 Word:{path}(段落 {n_para}、表格 {n_tbl})")
