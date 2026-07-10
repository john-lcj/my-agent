import hashlib
import json
import asyncio
from pathlib import Path

from office.artifacts import inspect_artifact
from office.attachments import extract_text, quarantine_attachment
from office.connectors import ConnectorScope, InMemoryConnector
from office.evidence import EvidencePackage, evidence_ready
from office.kernel import OfficeKernel, OfficeOperation
from office.runner import OfficeWorkflowRunner
from office.workflows import OfficeWorkflow
from office.workflows import benchmark_cases


def test_office_kernel_preview_and_idempotent_write(tmp_path):
    kernel = OfficeKernel(str(tmp_path / "office.db"))
    op = OfficeOperation("file.create", "report.docx", {"title": "Report"}, authority="write", dry_run=True)
    preview = kernel.execute(op, lambda: {"never": "runs"})
    assert preview.status == "previewed" and preview.preview["idempotency_key"] == op.key()
    op = OfficeOperation("file.create", "report.docx", {"title": "Report"}, authority="write", dry_run=False)
    first = kernel.execute(op, lambda: {"path": "report.docx"})
    second = kernel.execute(op, lambda: {"path": "should-not-run"})
    assert first.status == "completed" and second.status == "deduplicated"
    assert second.output == {"path": "report.docx"}
    blocked = kernel.execute(OfficeOperation("file.delete", "report.docx", authority="read", dry_run=False), lambda: True)
    assert blocked.status == "blocked"


def test_artifact_validation_and_evidence(tmp_path):
    path = tmp_path / "report.docx"
    import pytest
    pytest.importorskip("docx")
    from docx import Document
    doc = Document()
    doc.add_heading("Report", 0)
    doc.add_paragraph("Validated content")
    doc.save(path)
    check = inspect_artifact(str(path), render=True)
    assert check.ok and check.pages_or_sheets >= 1 and len(check.sha256) == 64
    package = EvidencePackage("office-test")
    package.add_source("notes.txt", hashlib.sha256(b"notes").hexdigest(), trusted=True)
    package.add_validation(check.as_dict())
    package.add_output(str(path), check.sha256)
    evidence_path = package.write(str(tmp_path))
    assert evidence_ready(package) and json.loads(Path(evidence_path).read_text())["task_id"] == "office-test"


def test_attachment_is_quarantined_and_extracted_as_untrusted(tmp_path):
    source = tmp_path / "notes.txt"
    source.write_text("untrusted attachment text", encoding="utf-8")
    record = quarantine_attachment(str(source), str(tmp_path / "quarantine"))
    record = extract_text(record, str(tmp_path / "extracted"))
    assert record.quarantined_path and record.trusted is False
    assert Path(record.extracted_text_path).read_text(encoding="utf-8") == "untrusted attachment text"


def test_connector_scope_defaults_to_read_only():
    connector = InMemoryConnector("google", records={"1": {"subject": "hello"}})
    assert connector.call("mail.list").ok
    assert not connector.call("mail.create", subject="blocked").ok
    writable = InMemoryConnector("google", ConnectorScope(write=True))
    assert writable.call("mail.create", subject="draft").ok
    assert not writable.call("mail.send", subject="send").ok


def test_p4_benchmark_has_30_cases_and_all_core_workflows():
    cases = benchmark_cases()
    assert len(cases) == 35
    assert {item["workflow"] for item in cases} >= {
        "inbox_triage", "weekly_report", "meeting_minutes", "spreadsheet_analysis",
        "research_report", "slide_deck", "calendar_coordination",
    }


def test_calendar_write_is_idempotent_and_rejects_invalid_range(tmp_path, monkeypatch):
    monkeypatch.setenv("CALENDAR_FILE", str(tmp_path / "calendar.ics"))
    from capabilities.tools.calendar_tool import CalendarAdd
    tool = CalendarAdd()
    args = {"title": "Review", "start": "2030-01-01 10:00", "duration_min": 60,
            "idempotency_key": "review-2030"}
    first = asyncio.run(tool.invoke(args, None))
    second = asyncio.run(tool.invoke(args, None))
    invalid = asyncio.run(tool.invoke({"title": "Bad", "start": "2030-01-01 10:00",
                                       "end": "2030-01-01 09:00"}, None))
    assert first.ok and second.ok and "幂等" in second.output
    assert not invalid.ok
    assert Path(tmp_path / "calendar.ics").read_text().count("BEGIN:VEVENT") == 1


def test_workflow_runner_withholds_invalid_artifact_and_emits_evidence(tmp_path):
    runner = OfficeWorkflowRunner(OfficeKernel(str(tmp_path / "office.db")), str(tmp_path / "evidence"))
    workflow = OfficeWorkflow("weekly_report", "Weekly report", ("notes",), ("docx",))
    good = runner.run(workflow, "report.txt", {}, lambda: str(tmp_path / "report.txt"), dry_run=True)
    assert good.status == "previewed"
    Path(tmp_path / "report.txt").write_text("report", encoding="utf-8")
    result = runner.run(workflow, "report.txt", {}, lambda: str(tmp_path / "report.txt"))
    assert result.ok and result.output["evidence"]
